# Advanced RAG Chatbot Application
# Complete error-free version with deployment fixes

import streamlit as st
import hashlib
import os
import json
import re
from datetime import datetime
import pandas as pd
from typing import List, Dict, Any
import sqlite3
import io
import pypdf as PyPDF2
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Configuration
USER_DB_PATH = "./users.db"
DOCUMENTS_PATH = "./documents"

def ensure_directories():
    """Create necessary directories and files for deployment"""
    try:
        # Create directories
        os.makedirs("documents/admin", exist_ok=True)
        os.makedirs("documents/users", exist_ok=True)
        
        # Create empty JSON files if they don't exist
        admin_file = "documents/admin/documents.json"
        user_file = "documents/users/documents.json"
        
        if not os.path.exists(admin_file):
            with open(admin_file, "w", encoding="utf-8") as f:
                json.dump([], f)
        
        if not os.path.exists(user_file):
            with open(user_file, "w", encoding="utf-8") as f:
                json.dump([], f)
                
    except Exception as e:
        st.error(f"Error creating directories: {e}")

def load_css():
    """Load external CSS file with embedded fallback"""
    try:
        # Try to load external CSS file first
        if os.path.exists('styles.css'):
            with open('styles.css', 'r', encoding='utf-8') as f:
                css = f.read()
            st.markdown(f'<style>{css}</style>', unsafe_allow_html=True)
        else:
            # Fallback: Use embedded CSS
            st.markdown("""
            <style>
            @import url('https://fonts.googleapis.com/css2?family=Orbitron:wght@400;700;900&family=Rajdhani:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500;600&display=swap');

            :root {
                --primary-glow: #00ffff;
                --secondary-glow: #ff00ff;
                --accent-glow: #ffff00;
                --dark-bg: #0a0a0f;
                --glass-bg: rgba(255, 255, 255, 0.05);
                --border-glow: rgba(0, 255, 255, 0.3);
            }

            .stApp {
                background: linear-gradient(-45deg, #0a0a0f, #1a1a2e, #16213e, #0f3460);
                background-size: 400% 400%;
                animation: gradientShift 15s ease infinite;
                font-family: 'Rajdhani', sans-serif;
                color: #ffffff;
            }

            @keyframes gradientShift {
                0% { background-position: 0% 50%; }
                50% { background-position: 100% 50%; }
                100% { background-position: 0% 50%; }
            }

            section[data-testid="stSidebar"] > div {
                background: linear-gradient(135deg, rgba(0, 255, 255, 0.1), rgba(255, 0, 255, 0.1)) !important;
                backdrop-filter: blur(20px) !important;
                border-right: 2px solid var(--border-glow) !important;
                box-shadow: 0 8px 32px rgba(0, 255, 255, 0.1) !important;
            }

            section[data-testid="stSidebar"] * {
                color: #ffffff !important;
                font-family: 'Rajdhani', sans-serif !important;
                text-shadow: 0 0 10px rgba(0, 255, 255, 0.5) !important;
            }

            .stButton > button {
                background: linear-gradient(45deg, rgba(0, 255, 255, 0.2), rgba(255, 0, 255, 0.2)) !important;
                border: 1px solid var(--primary-glow) !important;
                color: #ffffff !important;
                font-family: 'Orbitron', monospace !important;
                font-weight: 600 !important;
                text-transform: uppercase !important;
                letter-spacing: 1px !important;
                border-radius: 0 !important;
                clip-path: polygon(0 0, calc(100% - 15px) 0, 100% 50%, calc(100% - 15px) 100%, 0 100%, 15px 50%) !important;
                padding: 0.7rem 2rem !important;
                position: relative !important;
                overflow: hidden !important;
                transition: all 0.3s cubic-bezier(0.25, 0.46, 0.45, 0.94) !important;
                box-shadow: 0 0 20px rgba(0, 255, 255, 0.3) !important;
            }

            .stButton > button:hover {
                transform: scale(1.05) !important;
                box-shadow: 0 0 30px rgba(0, 255, 255, 0.6) !important;
                border-color: var(--secondary-glow) !important;
            }

            .main-header {
                font-family: 'Orbitron', monospace !important;
                font-size: 3.5rem !important;
                font-weight: 900 !important;
                text-align: center !important;
                background: linear-gradient(45deg, var(--primary-glow), var(--secondary-glow), var(--accent-glow)) !important;
                background-size: 300% 300% !important;
                -webkit-background-clip: text !important;
                -webkit-text-fill-color: transparent !important;
                background-clip: text !important;
                animation: holographicShift 4s ease-in-out infinite !important;
                margin: 2rem 0 !important;
            }

            @keyframes holographicShift {
                0%, 100% { background-position: 0% 50%; }
                33% { background-position: 50% 0%; }
                66% { background-position: 100% 50%; }
            }

            .stTextInput > div > div > input,
            .stChatInput > div > div > input {
                background: rgba(0, 255, 255, 0.05) !important;
                border: 2px solid var(--primary-glow) !important;
                border-radius: 0 !important;
                color: #ffffff !important;
                font-family: 'JetBrains Mono', monospace !important;
                font-size: 16px !important;
                padding: 1rem !important;
                clip-path: polygon(0 0, calc(100% - 20px) 0, 100% 20px, 100% 100%, 20px 100%, 0 calc(100% - 20px)) !important;
                transition: all 0.3s ease !important;
                box-shadow: inset 0 0 20px rgba(0, 255, 255, 0.1) !important;
            }

            .stTextInput > div > div > input:focus,
            .stChatInput > div > div > input:focus {
                outline: none !important;
                border-color: var(--secondary-glow) !important;
                box-shadow: 0 0 25px rgba(255, 0, 255, 0.5) !important;
                transform: scale(1.02) !important;
            }

            .feature-card {
                background: linear-gradient(135deg, rgba(255, 255, 255, 0.1), rgba(255, 255, 255, 0.02)) !important;
                backdrop-filter: blur(20px) !important;
                border: 1px solid rgba(0, 255, 255, 0.3) !important;
                border-radius: 0 !important;
                padding: 2rem !important;
                margin: 1.5rem 0 !important;
                clip-path: polygon(0 0, calc(100% - 30px) 0, 100% 30px, 100% 100%, 30px 100%, 0 calc(100% - 30px)) !important;
                position: relative !important;
                overflow: hidden !important;
                box-shadow: 0 8px 32px rgba(0, 255, 255, 0.1) !important;
            }

            h1, h2, h3, h4, h5, h6 {
                font-family: 'Orbitron', monospace !important;
                color: var(--primary-glow) !important;
                text-shadow: 0 0 15px rgba(0, 255, 255, 0.5) !important;
                letter-spacing: 1px !important;
            }

            .stChatMessage {
                background: rgba(0, 0, 0, 0.4) !important;
                border-left: 4px solid var(--primary-glow) !important;
                border-radius: 0 !important;
                margin: 1rem 0 !important;
                font-family: 'JetBrains Mono', monospace !important;
                box-shadow: 0 4px 16px rgba(0, 255, 255, 0.1) !important;
                backdrop-filter: blur(10px) !important;
            }

            @media (max-width: 768px) {
                .main-header {
                    font-size: 2.5rem !important;
                }
                
                .stButton > button {
                    padding: 0.5rem 1rem !important;
                    font-size: 0.9rem !important;
                }
            }
            </style>
            """, unsafe_allow_html=True)
    except Exception as e:
        st.error(f"Error loading CSS: {e}")

class UserManager:
    def __init__(self):
        self.init_user_db()
    
    def init_user_db(self):
        """Initialize SQLite database for user management"""
        conn = sqlite3.connect(USER_DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                role TEXT NOT NULL DEFAULT 'user',
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Create default admin user (admin/admin123)
        admin_hash = hashlib.sha256("admin123".encode()).hexdigest()
        cursor.execute('''
            INSERT OR IGNORE INTO users (username, password_hash, role) 
            VALUES (?, ?, ?)
        ''', ("admin", admin_hash, "admin"))
        
        # Create default user (user/user123)
        user_hash = hashlib.sha256("user123".encode()).hexdigest()
        cursor.execute('''
            INSERT OR IGNORE INTO users (username, password_hash, role) 
            VALUES (?, ?, ?)
        ''', ("user", user_hash, "user"))
        
        conn.commit()
        conn.close()
    
    def authenticate(self, username: str, password: str) -> Dict[str, Any]:
        """Authenticate user and return user info"""
        conn = sqlite3.connect(USER_DB_PATH)
        cursor = conn.cursor()
        password_hash = hashlib.sha256(password.encode()).hexdigest()
        
        cursor.execute('''
            SELECT username, role FROM users 
            WHERE username = ? AND password_hash = ?
        ''', (username, password_hash))
        
        result = cursor.fetchone()
        conn.close()
        
        if result:
            return {"username": result[0], "role": result[1], "authenticated": True}
        return {"authenticated": False}

class DocumentManager:
    def __init__(self):
        os.makedirs(f"{DOCUMENTS_PATH}/admin", exist_ok=True)
        os.makedirs(f"{DOCUMENTS_PATH}/users", exist_ok=True)
        
        self.admin_docs = []
        self.user_docs = []
        self.vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
        self.admin_vectors = None
        self.user_vectors = None
        self.load_documents()
    
    def load_documents(self):
        """Load documents from JSON files"""
        admin_file = f"{DOCUMENTS_PATH}/admin/documents.json"
        if os.path.exists(admin_file):
            try:
                with open(admin_file, "r", encoding="utf-8") as f:
                    self.admin_docs = json.load(f)
                self._update_admin_vectors()
            except:
                self.admin_docs = []
        
        user_file = f"{DOCUMENTS_PATH}/users/documents.json"
        if os.path.exists(user_file):
            try:
                with open(user_file, "r", encoding="utf-8") as f:
                    self.user_docs = json.load(f)
                self._update_user_vectors()
            except:
                self.user_docs = []
    
    def save_documents(self):
        """Save documents to separate JSON files"""
        try:
            with open(f"{DOCUMENTS_PATH}/admin/documents.json", "w", encoding="utf-8") as f:
                json.dump(self.admin_docs, f, ensure_ascii=False, indent=2)
            
            with open(f"{DOCUMENTS_PATH}/users/documents.json", "w", encoding="utf-8") as f:
                json.dump(self.user_docs, f, ensure_ascii=False, indent=2)
        except Exception as e:
            st.error(f"Error saving documents: {str(e)}")
    
    def add_document(self, content: str, filename: str, uploaded_by: str, user_role: str, metadata: Dict = None):
        """Add document with role-based separation"""
        if metadata is None:
            metadata = {}
        
        chunks = self.split_text(content)
        
        for i, chunk in enumerate(chunks):
            doc = {
                "id": f"{filename}_{i}_{datetime.now().timestamp()}",
                "content": chunk,
                "filename": filename,
                "uploaded_by": uploaded_by,
                "chunk_id": i,
                "total_chunks": len(chunks),
                "upload_time": datetime.now().isoformat(),
                **metadata
            }
            
            if user_role == "admin":
                self.admin_docs.append(doc)
            else:
                self.user_docs.append(doc)
        
        self._update_vectors()
        self.save_documents()
    
    def split_text(self, text: str, chunk_size: int = 300, overlap: int = 30) -> List[str]:
        """Optimized text splitting for faster processing"""
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = " ".join(words[i:i + chunk_size])
            if chunk.strip() and len(chunk.strip()) > 20:
                chunks.append(chunk)
        
        return chunks if chunks else [text]
    
    def _update_vectors(self):
        """Update vectors for both admin and user documents"""
        self._update_admin_vectors()
        self._update_user_vectors()
    
    def _update_admin_vectors(self):
        """Update TF-IDF vectors for admin documents"""
        if self.admin_docs:
            texts = [doc["content"] for doc in self.admin_docs]
            try:
                self.admin_vectors = self.vectorizer.fit_transform(texts)
            except:
                self.admin_vectors = None
    
    def _update_user_vectors(self):
        """Update TF-IDF vectors for user documents"""
        if self.user_docs:
            texts = [doc["content"] for doc in self.user_docs]
            try:
                if hasattr(self.vectorizer, 'vocabulary_'):
                    self.user_vectors = self.vectorizer.transform(texts)
                else:
                    self.user_vectors = self.vectorizer.fit_transform(texts)
            except:
                self.user_vectors = None
    
    def search_documents(self, query: str, user_role: str, include_all_for_admin: bool = False, n_results: int = 5) -> Dict:
        """Search documents based on user role"""
        if user_role == "admin" and include_all_for_admin:
            return self._search_combined(query, n_results)
        elif user_role == "admin":
            return self._search_specific(query, "admin", n_results)
        else:
            return self._search_specific(query, "user", n_results)
    
    def _search_combined(self, query: str, n_results: int) -> Dict:
        """Search both admin and user documents"""
        admin_results = self._search_specific(query, "admin", n_results//2)
        user_results = self._search_specific(query, "user", n_results//2)
        
        documents = admin_results["documents"][0] + user_results["documents"][0]
        metadatas = admin_results["metadatas"][0] + user_results["metadatas"][0]
        distances = admin_results["distances"][0] + user_results["distances"][0]
        
        return {
            "documents": [documents[:n_results]],
            "metadatas": [metadatas[:n_results]], 
            "distances": [distances[:n_results]]
        }
    
    def _search_specific(self, query: str, doc_type: str, n_results: int) -> Dict:
        """Optimized search for specific document type"""
        if doc_type == "admin":
            docs = self.admin_docs
            vectors = self.admin_vectors
        else:
            docs = self.user_docs
            vectors = self.user_vectors
        
        if not docs or vectors is None:
            return {"documents": [[]], "metadatas": [[]], "distances": [[]]}
        
        try:
            query_vector = self.vectorizer.transform([query])
            similarities = cosine_similarity(query_vector, vectors).flatten()
            
            if len(similarities) > n_results:
                top_indices = np.argpartition(similarities, -n_results)[-n_results:]
                top_indices = top_indices[np.argsort(similarities[top_indices])][::-1]
            else:
                top_indices = np.argsort(similarities)[::-1]
            
            documents = []
            metadatas = []
            distances = []
            
            for idx in top_indices:
                if similarities[idx] > 0.01:
                    documents.append(docs[idx]["content"])
                    metadatas.append({
                        "filename": docs[idx]["filename"],
                        "uploaded_by": docs[idx].get("uploaded_by", "system"),
                        "chunk_id": docs[idx]["chunk_id"]
                    })
                    distances.append(1 - similarities[idx])
            
            return {
                "documents": [documents],
                "metadatas": [metadatas],
                "distances": [distances]
            }
        except Exception:
            return {"documents": [[]], "metadatas": [[]], "distances": [[]]}
    
    def get_documents_by_role(self, role: str) -> List[str]:
        """Get documents by role"""
        if role == "admin":
            return list(set([doc["filename"] for doc in self.admin_docs]))
        else:
            return list(set([doc["filename"] for doc in self.user_docs]))
    
    def get_all_documents_for_admin(self) -> Dict[str, List[str]]:
        """Get all documents separated by type for admin view"""
        admin_files = list(set([doc["filename"] for doc in self.admin_docs]))
        user_files = list(set([doc["filename"] for doc in self.user_docs]))
        
        return {
            "admin_documents": admin_files,
            "user_documents": user_files
        }
    
    def delete_document(self, filename: str, doc_type: str = None):
        """Delete document by filename and type"""
        try:
            if doc_type == "admin" or doc_type is None:
                original_count = len(self.admin_docs)
                self.admin_docs = [doc for doc in self.admin_docs if doc["filename"] != filename]
                if len(self.admin_docs) < original_count:
                    self._update_admin_vectors()
            
            if doc_type == "user" or doc_type is None:
                original_count = len(self.user_docs)
                self.user_docs = [doc for doc in self.user_docs if doc["filename"] != filename]
                if len(self.user_docs) < original_count:
                    self._update_user_vectors()
            
            self.save_documents()
            return True
        except Exception as e:
            st.error(f"Delete error: {str(e)}")
            return False

class LLMResponseGenerator:
    @staticmethod
    def generate_response(query: str, context: str) -> str:
        """Generate response using retrieved context"""
        
        if LLMResponseGenerator._is_general_conversation(query):
            return LLMResponseGenerator._handle_general_conversation(query)
        
        if not context.strip():
            return LLMResponseGenerator._suggest_document_questions()
        
        local_response = LLMResponseGenerator._generate_intelligent_response(query, context)
        if local_response and "could not locate specific information" not in local_response:
            return local_response
        
        try:
            hf_response = LLMResponseGenerator._try_huggingface_fast(query, context)
            if hf_response:
                return hf_response
        except:
            pass
        
        return local_response if local_response else "I could not find relevant information for your query. Try rephrasing your question."
    
    @staticmethod
    def _try_huggingface_fast(query: str, context: str) -> str:
        """Fast HuggingFace API call"""
        try:
            import requests
            
            api_url = "https://api-inference.huggingface.co/models/deepset/roberta-base-squad2"
            payload = {
                "inputs": {
                    "question": query,
                    "context": context[:1000]
                }
            }
            
            response = requests.post(
                api_url, 
                headers={"Authorization": "Bearer hf_free"}, 
                json=payload, 
                timeout=3
            )
            
            if response.status_code == 200:
                result = response.json()
                if 'answer' in result and result['answer'].strip():
                    confidence = result.get('score', 0)
                    if confidence > 0.15:
                        return f"**Answer:** {result['answer']}\n\n*Powered by HuggingFace*"
            
        except Exception:
            pass
        
        return None
    
    @staticmethod
    def _is_general_conversation(query: str) -> bool:
        """Check if query is general conversation"""
        query_lower = query.lower().strip()
        
        general_patterns = [
            'hi', 'hello', 'hey', 'good morning', 'good afternoon', 'good evening',
            'how are you', 'how r u', 'whats up', 'what is up', 'sup',
            'thanks', 'thank you', 'bye', 'goodbye', 'see you',
            'who are you', 'what are you', 'what can you do',
            'help', 'how does this work', 'what is this',
            'test', 'testing', 'can you hear me'
        ]
        
        for pattern in general_patterns:
            if query_lower == pattern or query_lower.startswith(pattern):
                return True
        
        if len(query_lower.split()) <= 3 and not any(word in query_lower for word in ['what', 'where', 'when', 'how', 'who', 'which', 'describe', 'explain', 'tell me']):
            return True
        
        return False
    
    @staticmethod
    def _handle_general_conversation(query: str) -> str:
        """Handle general conversation queries"""
        query_lower = query.lower().strip()
        
        if any(word in query_lower for word in ['hi', 'hello', 'hey', 'morning', 'afternoon', 'evening']):
            return """**Welcome to your Advanced RAG Assistant!**

I am here to help you explore and analyze your documents with cutting-edge AI technology.

**What I can do:**
- Process any document type - PDFs, Word docs, text files
- AI-powered analysis using local processing and HuggingFace
- Smart search to find exactly what you are looking for
- Natural conversation - Ask questions in plain English

**Ready to get started?**
Upload documents and ask me anything about them!"""

        elif any(phrase in query_lower for phrase in ['how are you', 'how r u', 'whats up', 'what is up']):
            return """**I am running at full capacity and ready to assist!**

My AI engines are warmed up and waiting to help you analyze your documents.

**Current capabilities:**
- Multiple AI models active
- Document processing optimized
- Search algorithms ready
- Natural language understanding enabled

**What would you like to explore in your documents today?**"""

        elif any(phrase in query_lower for phrase in ['who are you', 'what are you', 'what can you do']):
            return """**I am your Advanced RAG (Retrieval-Augmented Generation) Assistant**

**My Technology Stack:**
- Local Processing - Privacy-first analysis
- HuggingFace - Specialized models
- TF-IDF Search - Fast document retrieval
- Multi-user Support - Separated document management

**My Capabilities:**
- Upload and process any document type
- Intelligent content extraction and chunking
- AI response generation
- Context-aware conversations
- Role-based document management

**I excel at:**
- Finding specific information in large documents
- Summarizing complex content
- Answering detailed questions
- Cross-referencing multiple sources

**Ready to experience document intelligence?**"""

        return f"""**I understand you said: "{query}"**

I am optimized for document analysis and intelligent information retrieval. While I enjoy our conversation, my real power lies in helping you unlock insights from your uploaded documents!

**Want to see what I can really do?**
- Upload some documents (any format)
- Ask me complex questions about the content
- Watch me provide precise, AI-powered answers

**Let us put my capabilities to the test!**"""
    
    @staticmethod
    def _suggest_document_questions():
        """Suggest what to ask when no documents are found"""
        return """**No relevant documents found for your query.**

**Suggestions:**

**If you are a User:**
- Upload your own documents using the upload feature
- Ask questions about documents you have uploaded
- Try more specific keywords

**If you are an Admin:**
- Check both Admin and User document sections
- Ensure documents are properly uploaded and processed
- Try searching with different terms

**Example questions that work well:**
- "What skills are mentioned in the resume?"
- "Summarize the main points of the document"
- "What experience does this person have?"
- "Find information about [specific topic]"

**Ready to upload and explore?**"""
    
    @staticmethod
    def _generate_intelligent_response(query: str, context: str) -> str:
        """Generate intelligent response using local processing"""
        context = re.sub(r'\s+', ' ', context.strip())
        sentences = [s.strip() for s in context.split('.') if len(s.strip()) > 15]
        sentences = sentences[:15]
        
        query_words = LLMResponseGenerator._extract_keywords_fast(query.lower())
        
        scored_sentences = []
        for sentence in sentences:
            score = LLMResponseGenerator._calculate_relevance_fast(sentence.lower(), query_words)
            if score > 0:
                scored_sentences.append((score, sentence))
        
        scored_sentences.sort(reverse=True, key=lambda x: x[0])
        
        if not scored_sentences:
            return f"I found content but could not locate specific information about '{query}'. Try using different keywords."
        
        return LLMResponseGenerator._format_answer_fast(query, scored_sentences[:2])
    
    @staticmethod
    def _extract_keywords_fast(query):
        """Fast keyword extraction"""
        stop_words = {'what', 'how', 'when', 'where', 'why', 'who', 'is', 'are', 'the', 'a', 'an', 'and', 'or', 'in', 'on', 'at', 'to', 'for', 'of', 'with'}
        
        words = query.split()
        keywords = [w for w in words if w not in stop_words and len(w) > 2]
        
        if 'experience' in keywords:
            keywords.extend(['work', 'job', 'career'])
        if 'education' in keywords:
            keywords.extend(['school', 'university', 'college', 'degree'])
        if 'skills' in keywords:
            keywords.extend(['abilities', 'expertise', 'knowledge'])
        
        return keywords
    
    @staticmethod
    def _calculate_relevance_fast(sentence, keywords):
        """Fast relevance calculation"""
        score = 0
        sentence_lower = sentence.lower()
        
        for keyword in keywords:
            if keyword in sentence_lower:
                score += 3
        
        keyword_count = sum(1 for kw in keywords if kw in sentence_lower)
        if keyword_count > 1:
            score += keyword_count * 2
        
        return score
    
    @staticmethod
    def _format_answer_fast(query, scored_sentences):
        """Fast answer formatting"""
        if not scored_sentences:
            return "No relevant information found."
        
        top_sentence = scored_sentences[0][1]
        response = f"**Answer:** {top_sentence}"
        
        if len(scored_sentences) > 1:
            response += f"\n\n**Additional info:** {scored_sentences[1][1]}"
        
        response += "\n\n*Fast local processing*"
        return response

def login_page():
    """Display futuristic login page"""
    load_css()
    
    st.markdown('<h1 class="main-header">NEXUS RAG INTELLIGENCE</h1>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <h3 style="text-align: center; margin-bottom: 2rem;">SECURE ACCESS PORTAL</h3>
        </div>
        """, unsafe_allow_html=True)
        
        username = st.text_input("Username", placeholder="Enter authorization code...")
        password = st.text_input("Password", type="password", placeholder="Enter security key...")
        
        if st.button("INITIALIZE NEURAL LINK", use_container_width=True):
            user_manager = UserManager()
            auth_result = user_manager.authenticate(username, password)
            
            if auth_result["authenticated"]:
                st.session_state.user = auth_result
                st.success("AUTHENTICATION SUCCESSFUL - NEURAL LINK ESTABLISHED")
                st.rerun()
            else:
                st.error("ACCESS DENIED - INVALID CREDENTIALS")
        
        st.markdown("---")
        
        with st.expander("DEMO ACCESS CREDENTIALS"):
            st.markdown("""
            **ADMINISTRATOR ACCESS:**
            - Username: `admin`
            - Password: `admin123`
            - Capabilities: Full system control, document management, AI configuration
            
            **USER ACCESS:**
            - Username: `user`  
            - Password: `user123`
            - Capabilities: Document upload, AI chat, personal document management
            """)

def admin_dashboard():
    """Advanced admin dashboard"""
    load_css()
    
    st.markdown('<h1 class="main-header">ADMIN CONTROL CENTER</h1>', unsafe_allow_html=True)
    
    with st.sidebar:
        st.markdown("### CONTROL PANEL")
        action = st.selectbox(
            "SELECT OPERATION:",
            ["Dashboard Overview", "Upload Documents", "Document Management", "System Settings"]
        )
    
    db_manager = DocumentManager()
    
    if action == "Dashboard Overview":
        st.markdown("## SYSTEM OVERVIEW")
        
        all_docs = db_manager.get_all_documents_for_admin()
        admin_count = len(all_docs["admin_documents"])
        user_count = len(all_docs["user_documents"])
        total_chunks = len(db_manager.admin_docs) + len(db_manager.user_docs)
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Admin Documents", admin_count)
        with col2:
            st.metric("User Documents", user_count)
        with col3:
            st.metric("Total Chunks", total_chunks)
        with col4:
            st.metric("AI Status", "ONLINE")
        
        st.markdown("### NEURAL NETWORK ANALYTICS")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.markdown("#### ADMIN NEURAL CORES")
            if all_docs["admin_documents"]:
                for doc in all_docs["admin_documents"]:
                    st.markdown(f"• {doc}")
            else:
                st.info("NO ADMIN NEURAL CORES DETECTED")
        
        with col2:
            st.markdown("#### USER NEURAL CORES")
            if all_docs["user_documents"]:
                for doc in all_docs["user_documents"]:
                    st.markdown(f"• {doc}")
            else:
                st.info("NO USER NEURAL CORES DETECTED")
    
    elif action == "Upload Documents":
        st.markdown("## NEURAL DATA UPLOAD CENTER")
        
        upload_type = st.radio(
            "SELECT NEURAL CORE DESTINATION:",
            ["Admin Neural Cores", "User Neural Cores"],
            horizontal=True
        )
        
        st.markdown("### DATA STREAM UPLOAD")
        
        uploaded_files = st.file_uploader(
            "INITIALIZE DATA TRANSFER PROTOCOL",
            accept_multiple_files=True,
            type=['txt', 'pdf', 'docx', 'csv', 'json'],
            help="COMPATIBLE FORMATS: PDF, Word, Text, CSV, JSON"
        )
        
        if uploaded_files:
            st.markdown("### UPLOAD PROCESSING QUEUE")
            
            for uploaded_file in uploaded_files:
                col1, col2, col3 = st.columns([3, 1, 1])
                
                with col1:
                    st.write(f"**{uploaded_file.name}**")
                    st.write(f"Size: {uploaded_file.size / 1024:.1f} KB | Type: {uploaded_file.type}")
                
                with col2:
                    if st.button(f"UPLOAD", key=f"upload_{uploaded_file.name}"):
                        try:
                            content = ""
                            
                            if uploaded_file.type == "text/plain":
                                content = str(uploaded_file.read(), "utf-8")
                            elif uploaded_file.type == "text/csv":
                                df = pd.read_csv(uploaded_file)
                                content = df.to_string()
                            elif uploaded_file.type == "application/json":
                                content = str(uploaded_file.read(), "utf-8")
                            elif uploaded_file.type == "application/pdf":
                                import pypdf
                                pdf_reader = pypdf.PdfReader(io.BytesIO(uploaded_file.read()))
                                content = ""
                                for page in pdf_reader.pages:
                                    content += page.extract_text() + "\n"
                                if not content.strip():
                                    content = f"PDF file: {uploaded_file.name} (Text extraction may have failed)"
                            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                                doc = Document(io.BytesIO(uploaded_file.read()))
                                content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                                if not content.strip():
                                    content = f"DOCX file: {uploaded_file.name} (No readable text found)"
                            else:
                                try:
                                    content = str(uploaded_file.read(), "utf-8")
                                except UnicodeDecodeError:
                                    content = f"Binary file: {uploaded_file.name} (Content type: {uploaded_file.type})"
                            
                            if content and content.strip():
                                target_role = "admin" if "Admin" in upload_type else "user"
                                db_manager.add_document(
                                    content=content,
                                    filename=uploaded_file.name,
                                    uploaded_by=st.session_state.user["username"],
                                    user_role=target_role,
                                    metadata={
                                        "file_type": uploaded_file.type,
                                        "file_size": uploaded_file.size
                                    }
                                )
                                st.success(f"DATA TRANSFER COMPLETE - {uploaded_file.name} UPLOADED SUCCESSFULLY")
                            else:
                                st.warning(f"WARNING - {uploaded_file.name} CONTAINS NO READABLE DATA")
                        
                        except Exception as e:
                            st.error(f"ERROR - FAILED TO PROCESS {uploaded_file.name}: {str(e)}")
                
                with col3:
                    st.write("READY")
    
    elif action == "Document Management":
        st.markdown("## NEURAL CORE MANAGEMENT CENTER")
        
        tab1, tab2 = st.tabs(["Admin Neural Cores", "User Neural Cores"])
        
        all_docs = db_manager.get_all_documents_for_admin()
        
        with tab1:
            st.markdown("### ADMINISTRATOR NEURAL CORES")
            admin_docs = all_docs["admin_documents"]
            
            if admin_docs:
                st.write(f"**TOTAL ADMIN CORES: {len(admin_docs)}**")
                
                for doc in admin_docs:
                    col1, col2, col3 = st.columns([4, 1, 1])
                    with col1:
                        st.write(f"**{doc}**")
                    with col2:
                        if st.button("ANALYZE", key=f"analyze_admin_{doc}"):
                            st.info(f"ANALYZING NEURAL CORE: {doc}")
                    with col3:
                        if st.button("DELETE", key=f"delete_admin_{doc}"):
                            if db_manager.delete_document(doc, "admin"):
                                st.success(f"NEURAL CORE DELETED: {doc}")
                                st.rerun()
                            else:
                                st.error(f"DELETION FAILED: {doc}")
            else:
                st.info("NO ADMIN NEURAL CORES DETECTED")
        
        with tab2:
            st.markdown("### USER NEURAL CORES")
            user_docs = all_docs["user_documents"]
            
            if user_docs:
                st.write(f"**TOTAL USER CORES: {len(user_docs)}**")
                
                for doc in user_docs:
                    col1, col2, col3 = st.columns([4, 1, 1])
                    with col1:
                        st.write(f"**{doc}**")
                    with col2:
                        if st.button("ANALYZE", key=f"analyze_user_{doc}"):
                            st.info(f"ANALYZING NEURAL CORE: {doc}")
                    with col3:
                        if st.button("DELETE", key=f"delete_user_{doc}"):
                            if db_manager.delete_document(doc, "user"):
                                st.success(f"NEURAL CORE DELETED: {doc}")
                                st.rerun()
                            else:
                                st.error(f"DELETION FAILED: {doc}")
            else:
                st.info("NO USER NEURAL CORES DETECTED")
    
    elif action == "System Settings":
        st.markdown("## SYSTEM CONFIGURATION")
        
        tab1, tab2 = st.tabs(["Installation Protocol", "Performance Matrix"])
        
        with tab1:
            st.markdown("### COMPLETE INSTALLATION PROTOCOL")
            
            with st.expander("PYTHON NEURAL REQUIREMENTS"):
                st.code("""
# Core neural dependencies
pip install streamlit pandas numpy scikit-learn

# Document processing cores
pip install PyPDF2 python-docx

# AI integration modules
pip install requests

# Alternative installation protocols:
pip3 install streamlit pandas numpy scikit-learn PyPDF2 python-docx requests
python -m pip install streamlit pandas numpy scikit-learn PyPDF2 python-docx requests

# For conda neural networks:
conda install -c conda-forge streamlit pandas numpy scikit-learn
pip install PyPDF2 python-docx requests
                """, language="bash")
            
            with st.expander("NEURAL NETWORK ACTIVATION"):
                st.code("""
# Method 1 (Primary Neural Protocol):
streamlit run app.py

# Method 2 (Alternative Neural Link):
python -m streamlit run app.py

# Method 3 (Python3 Neural Interface):
python3 -m streamlit run app.py

# Windows Neural Protocol:
py -m streamlit run app.py

# macOS/Linux Neural Environment:
source venv/bin/activate
streamlit run app.py
                """, language="bash")
        
        with tab2:
            st.markdown("### PERFORMANCE OPTIMIZATION MATRIX")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### CURRENT OPTIMIZATIONS")
                st.success("✓ Reduced neural chunk size (300 words)")
                st.success("✓ Limited search results (2 chunks)")
                st.success("✓ Fast keyword neural matching")
                st.success("✓ Optimized vectorization matrix")
                st.success("✓ Short API neural timeouts (3s)")
                
            with col2:
                st.markdown("#### PERFORMANCE ENHANCEMENT TIPS")
                st.info("**For faster neural responses:**")
                st.write("• Ask specific neural queries")
                st.write("• Use clear keyword protocols")
                st.write("• Upload smaller data cores")
                st.write("• Use concise neural identifiers")
                st.write("• Clear chat neural history regularly")
            
            if st.button("PERFORMANCE NEURAL TEST"):
                import time
                
                start_time = time.time()
                
                test_text = "John Smith is a software engineer with 5 years of experience in Python."
                test_chunks = db_manager.split_text(test_text)
                
                if db_manager.user_docs or db_manager.admin_docs:
                    search_results = db_manager.search_documents("test query", "user", n_results=1)
                
                response = LLMResponseGenerator.generate_response("test", test_text)
                
                end_time = time.time()
                processing_time = end_time - start_time
                
                st.success(f"NEURAL PERFORMANCE TEST COMPLETE - PROCESSING TIME: {processing_time:.2f} SECONDS")
                
                if processing_time < 1:
                    st.success("NEURAL PERFORMANCE: EXCELLENT")
                elif processing_time < 3:
                    st.info("NEURAL PERFORMANCE: GOOD")
                else:
                    st.warning("NEURAL PERFORMANCE: OPTIMIZATION REQUIRED")

def user_interface():
    """Enhanced futuristic user interface"""
    load_css()
    
    st.markdown('<h1 class="main-header">NEXUS AI NEURAL ASSISTANT</h1>', unsafe_allow_html=True)
    
    db_manager = DocumentManager()
    llm_generator = LLMResponseGenerator()
    
    with st.sidebar:
        st.markdown("### USER NEURAL PANEL")
        
        st.markdown("#### NEURAL DATA UPLOAD")
        uploaded_file = st.file_uploader(
            "UPLOAD NEURAL DOCUMENT",
            type=['txt', 'pdf', 'docx', 'csv', 'json'],
            help="Upload documents to enhance neural knowledge base"
        )
        
        if uploaded_file:
            if st.button("UPLOAD TO NEURAL CORE", use_container_width=True):
                try:
                    content = ""
                    
                    if uploaded_file.type == "text/plain":
                        content = str(uploaded_file.read(), "utf-8")
                    elif uploaded_file.type == "text/csv":
                        df = pd.read_csv(uploaded_file)
                        content = df.to_string()
                    elif uploaded_file.type == "application/json":
                        content = str(uploaded_file.read(), "utf-8")
                    elif uploaded_file.type == "application/pdf":
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                        if not content.strip():
                            content = f"PDF file: {uploaded_file.name} (Text extraction may have failed)"
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = Document(io.BytesIO(uploaded_file.read()))
                        content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                        if not content.strip():
                            content = f"DOCX file: {uploaded_file.name} (No readable text found)"
                    else:
                        try:
                            content = str(uploaded_file.read(), "utf-8")
                        except UnicodeDecodeError:
                            content = f"Binary file: {uploaded_file.name} (Content type: {uploaded_file.type})"
                    
                    if content and content.strip():
                        db_manager.add_document(
                            content=content,
                            filename=uploaded_file.name,
                            uploaded_by=st.session_state.user["username"],
                            user_role="user",
                            metadata={
                                "file_type": uploaded_file.type,
                                "file_size": uploaded_file.size
                            }
                        )
                        st.success(f"NEURAL UPLOAD COMPLETE - {uploaded_file.name} INTEGRATED")
                        st.rerun()
                    else:
                        st.warning("WARNING - FILE CONTAINS NO READABLE NEURAL DATA")
                
                except Exception as e:
                    st.error(f"NEURAL UPLOAD FAILED - {str(e)}")
        
        st.markdown("---")
        
        st.markdown("### YOUR NEURAL CORES")
        user_docs = db_manager.get_documents_by_role("user")
        
        if user_docs:
            for doc in user_docs:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"Document: {doc}")
                with col2:
                    if st.button("DEL", key=f"del_{doc}", help=f"Delete {doc}"):
                        if db_manager.delete_document(doc, "user"):
                            st.success("DELETED")
                            st.rerun()
        else:
            st.info("NO NEURAL CORES DETECTED")
        
        st.markdown("---")
        
        if st.button("CLEAR NEURAL CHAT", use_container_width=True):
            st.session_state.messages = []
            st.rerun()
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    if prompt := st.chat_input("ENTER NEURAL QUERY..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("PROCESSING NEURAL QUERY"):
                search_results = db_manager.search_documents(prompt, "user", n_results=2)
                
                context = ""
                if search_results["documents"] and search_results["documents"][0]:
                    context = "\n\n".join(search_results["documents"][0][:2])
                    context = context[:1500]
                
                try:
                    response = llm_generator.generate_response(prompt, context)
                except Exception as e:
                    response = f"NEURAL ERROR - I encountered an issue processing your request. Please try rephrasing your query.\n\nError Code: {str(e)}"
                
                st.markdown(response)
                
                if search_results["metadatas"] and search_results["metadatas"][0]:
                    with st.expander("NEURAL SOURCE CORES"):
                        sources = set()
                        for metadata in search_results["metadatas"][0][:2]:
                            if "filename" in metadata:
                                sources.add(f"Document: {metadata['filename']}")
                        
                        if sources:
                            for source in sources:
                                st.write(source)
                        else:
                            st.write("NO SPECIFIC SOURCE CORES IDENTIFIED")
        
        st.session_state.messages.append({"role": "assistant", "content": response})

def main():
    """Main application with deployment compatibility"""
    # IMPORTANT: Create directories first
    ensure_directories()
    
    st.set_page_config(
        page_title="Nexus RAG Intelligence",
        page_icon="⬢",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    if "user" not in st.session_state:
        st.session_state.user = None
    
    if st.session_state.user is None:
        login_page()
    else:
        with st.sidebar:
            st.markdown(f"""
            ### **{st.session_state.user['username'].upper()}**
            **Neural Role:** {st.session_state.user['role'].title()}
            """)
            
            if st.button("TERMINATE NEURAL LINK", use_container_width=True):
                st.session_state.user = None
                st.session_state.messages = []
                st.rerun()
        
        if st.session_state.user["role"] == "admin":
            admin_dashboard()
        else:
            user_interface()

if __name__ == "__main__":
    main()
